#!/usr/bin/python3
#==============================================================================
# This is a simple generic implementation of a basic web harvester using
# BeautifulSoup combined with the Excel(openpyxl) and Word(docx) tools for Python. 
# It must be properly customized before use. Use at your own risk.
#==============================================================================

# Imports
import os
import re
import openpyxl
import requests
from requests import adapters
import string
from docx import Document
from docx.enum.text import WD_BREAK
from bs4 import BeautifulSoup
# End


def extractor(): # Extract url strings from Excel file

	wb = openpyxl.load_workbook('MyWorkbook.xlsx') # Set up openpyxl (Excel package for Python)
	sheet = wb['SheetName']
	index = 0
	count = 0
	for index in range(0,0,0): # Iterate over cells
		flag = 0
		print ('\n\n\n\n')
		print ('ITERATION:', index) # Print iteration in terminal
		url = sheet['B%s' %index].value # Extract url
		title = sheet['A%s' %index].value # Extract title
		print (title) # Print title in terminal
		if url != None: # Verify url is valid
			s = requests.Session() # Send request to server
			a = requests.adapters.HTTPAdapter(max_retries=5) # Configure retries
			s.mount('http://', a)
			response = s.get(url)
			html = response.text # Convert response format
			soup = BeautifulSoup(html, "html.parser") # Initiate BeautifulSoup package for Python
			text = soup.get_text() # Get text from raw data
			valid = bool(re.search('(?i)MyData ', title)) # Verify data is valid by title name
			if valid == True:
				print ('IS VALID')
				word_builder(soup, url, title) # Send raw data to be cleaned
	return


def word_builder(soup, url, title): # Extract relevant content and clean response 

	# Remove html style elements
	for script in source_soup(["script", "style"]):
		script.extract()
	soup.prettify()

	# Locate specific content within response in various ways
	temp_text = soup.find_all(True, class_="ClassName") # By class name
	if not temp_text:
			temp_text = soup.find_all(id=['IDname']) # By ID name
			if not temp_text:
				temp_text = soup.find_all(text=re.compile('(?i)ContentName')) # By content name
				flag = 1
				if not temp_text:
					return

 	# Some more clean up (Lines 72-76 based on StackOverflow post) 
	text = ''
	if not flag:
		for objects in temp_text: 
			text += objects.get_text(separator=u'\n')
		# Remove spaces
		lines = (line.strip() for line in text.splitlines())
		# Reorder in line format
		chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
		# Remove blank lines
		text = ' '.join(chunk for chunk in chunks if chunk)

	word_function(title, url, text)


# Build word document
def word_function(title, url, text):

	doc = Document()
	doc.add_heading(title)
	doc.add_paragraph('\n')
	doc.add_paragraph(url)
	doc.add_paragraph(text)
	doc.add_paragraph('\n')

	# Clean filename before saving
	clean_filename =  re.sub('[\\\\\r\n/:*?â€“"<>|]', '', title)
	doc.save(clean_filename[:173] + '.docx')

	print ('DOCUMENT IS DONE\n\n\n\n')
	return

# THE MAIN FUNCTION=========================================

extractor()

# ============================================================

